from __future__ import division
from collections import defaultdict
from collections import OrderedDict
from docx import Document
import decimal
import math
from random import randint
import calendar
import xlrd
import xlsxwriter
import sys
import os



months=["January","February","March","April","May","June","July"
		,"August","September","October","November","December"]	


def leading_Zero(num):
	num=str(num)
	if len(num)==1:
		num="0"+num
	return num

def in_dict_abbrev(key,out_list):
	key=key.split()
	for name in out_list:
		check=name.split()
		if (key[0]==check[0]):
			if (len(check)>1):
				if key[1][:1]==check[1]:
					return True
				else:
					continue
			else:
				return True
		else:
			continue
	return False




def date_help(year,month):
	date=calendar.monthrange(year,month)
	monthOne_Days=date[1]
	year_2=year
	if (month==12):
		year_2=year+1
		date_2=calendar.monthrange(year_2,1)
	else:
		date_2=calendar.monthrange(year,month+1)
	monthTwo_Days=date_2[1]
	return monthOne_Days,monthTwo_Days,year_2

def parse_cods(cod_levels_file):
	codL1=defaultdict(int)
	codL2=defaultdict(int)
	codL3=defaultdict(int)
	codL4=defaultdict(int)

	workbook = xlrd.open_workbook(cod_levels_file)
	sheet=workbook.sheet_by_index(0)
	rows=sheet.nrows
	cols=sheet.ncols
	data = [[sheet.cell_value(r,c) for c in range(sheet.ncols)] for r in range(sheet.nrows)]
	for row in range(1,rows):
		if sheet.cell_value(row,1)==1:
			codL1[sheet.cell_value(row,0).lower()]=0
		if sheet.cell_value(row,1)==2:
			codL2[sheet.cell_value(row,0).lower()]=0
		if sheet.cell_value(row,1)==3:
			codL3[sheet.cell_value(row,0).lower()]=0
		if sheet.cell_value(row,1)==4:
			codL4[sheet.cell_value(row,0).lower()]=0




	return codL1,codL2,codL3,codL4

def shift_Alloc(c1,c2,c3,c4,total_Shifts):
	len_C1=len(c1)-1
	len_C2=len(c2)
	len_C3=len(c3)
	len_C4=len(c4)
	S_4=(total_Shifts/(10*len_C1))
	level_1=math.ceil(4*S_4)
	level_2=math.ceil(3*S_4)
	level_3=math.ceil(2*S_4)
	level_4=math.ceil(S_4)
	level_Shifts=[level_1,level_2,level_3,level_4]
	return level_Shifts


def parse_travel(cod_out_file,year,month):
	dateDict_1=defaultdict(list)
	dateDict_2=defaultdict(list)
	m1_num,m2_num,year_2=date_help(year,month)
	mid=0;

	for l in open(cod_out_file,"r+"):
		l=l.split()
		if "DTEND;VALUE" in l[0]:
			end_Date=l[0].split(":")[1][4:]
		if "DTEND;TZID" in l[0]:
			mid=1
			end_Date=l[2].split(":")[1]
			end_Date=end_Date[4:]
			end_Date=end_Date.split("T")[0]
		if "DTSTART;VALUE" in l[0]:
			start_Date=l[0].split(":")[1][4:]
		if "DTSTART;TZID" in l[0]:
			start_Date=l[2].split(":")[1]
			start_Date=start_Date[4:]
			start_Date=start_Date.split("T")[0]
		if "SUMMARY" in l[0]:
			cod_Name=l[0].split(":")[1]
			if (cod_Name==""):
				cod_Name=l[1]
			cod_Name=cod_Name.lower()

		
			if (len(l)>1):
				if ((len(l[1])<3) and l[1].isupper() and (l[1]!="NY")):
					cod_Name+=" {0}".format(l[1][:1])
					cod_Name=cod_Name.lower()
			if (end_Date[:2]!=start_Date[:2]):
				for i in range(int(float(start_Date[2:])),m1_num):
					dateDict_1[start_Date[:2]+"/"+leading_Zero(i)].append(cod_Name.lower())
				for j in range(1,int(float(end_Date[2:]))):
					dateDict_2[end_Date[:2]+"/"+leading_Zero(j)].append(cod_Name.lower())
			else:
				if (start_Date[:2]==leading_Zero(month)):
					for i in range(int(start_Date[2:]),int(end_Date[2:])):
						dateDict_1[start_Date[:2]+"/"+leading_Zero(i)].append(cod_Name.lower())
				elif ((start_Date[:2])==leading_Zero(month+1)):
						for j in range(int(float(start_Date[2:])),int(end_Date[2:])):
							dateDict_2[start_Date[:2]+"/"+leading_Zero(j)].append(cod_Name.lower())
		
								
	return dateDict_1,dateDict_2	

				

def dict_random(dic):
	for key,val in dic.iteritems():
		flag=randint(0,1)
		if flag:
			del dic[key]
			dic[key]=val

		

		

def COD_Scheduler_3000(in_file,out_file,month,year,total_shifts):
	days=["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]
	not_out=defaultdict(int)
	work_Heap=OrderedDict()
	working=[]
	three_Shifts=0
	c1=defaultdict()
	c2=defaultdict()
	c3=defaultdict()
	c4=defaultdict()
	if (month==12):
		month_2=1
	else:
		month_2=int(month)+1
	monthOne_Days,monthTwo_Days,year_2=date_help(year,month)
	c1,c2,c3,c4=parse_cods(in_file)
	cod_Dict_List=[c1,c2,c3,c4]
	for i in range(4):
		dic=cod_Dict_List[i]
		for key,val in dic.iteritems():
			work_Heap[key]="{0}:0".format(str(i),str(0))
	dict_random(work_Heap)
	#Initialized 4 dicts with CODs by level, shifts start at 0
	shift_Limit=shift_Alloc(c1,c2,c3,c4,total_shifts)
	month_1_Out, month_2_Out = parse_travel(out_file,year,month)
	# initialized 2 dicts with key as date and value as lift of staff out
	cod_Dict_List=[c1,c2,c3,c4]
	document = Document()
	document.add_heading('Counselors In The Office', 0)
	workbook = xlsxwriter.Workbook('COD_Schedule.xlsx')
	worksheet = workbook.add_worksheet(months[month-1])
	row=0
	row_inc=0
	col=1
	for d in range (1,monthOne_Days+1):
		date=leading_Zero(month)+"/"+leading_Zero(d)
		day=calendar.weekday(year,month,d)
		row=row_inc
		if(day!=6):
			for key,val in work_Heap.iteritems():
				val_split=val.split(":")
				key_Level=val_split[0]
				key_Shifts=val_split[1]
				if not in_dict_abbrev(key,month_1_Out[date]):
					not_out[key]=val_split[1]
					if (three_Shifts!=3):
						if int(key_Shifts)!=shift_Limit[int(key_Level)]:
							three_Shifts+=1
							working.append(key)
							del work_Heap[key]
							work_Heap[key]=val_split[0]+":"+str(int(val_split[1])+1)
		
			worksheet.write(row,col,(days[day]+"  "+date))
			for cod in working:
				row+=1
				worksheet.write(row,col,cod)
			paragraph = document.add_paragraph()
			run = paragraph.add_run(days[day]+"  "+date+"\n",'Emphasis')
			for key, value in not_out.iteritems():
				paragraph.add_run(key+" has worked "+str(value)+" shifts\n")
			three_Shifts=0
			del working[:]
			not_out.clear()
			col+=2
		else:
			col=1
			row_inc+=5


	row=0
	row_inc=0
	col=1		
	worksheet = workbook.add_worksheet(months[month_2-1])
	for d in range (1,monthTwo_Days+1):
		date=leading_Zero(month_2)+"/"+leading_Zero(d)
		day=calendar.weekday(year,month_2,d)
		row=row_inc
		if(day!=6):
			for key,val in work_Heap.iteritems():
				val_split=val.split(":")
				key_Level=val_split[0]
				key_Shifts=val_split[1]
				if not in_dict_abbrev(key,month_2_Out[date]):
					not_out[key]=val_split[1]
					if (three_Shifts!=3):
						if int(key_Shifts)!=shift_Limit[int(key_Level)]:
							three_Shifts+=1
							working.append(key)
							del work_Heap[key]
							work_Heap[key]=val_split[0]+":"+str(int(val_split[1])+1)
								
			worksheet.write(row,col,(days[day]+"  "+date))
			for cod in working:
				row+=1
				worksheet.write(row,col,cod)
			paragraph = document.add_paragraph()
			run = paragraph.add_run(days[day]+"  "+date+"\n",'Emphasis')
			for key, value in not_out.iteritems():
				paragraph.add_run(key+" has worked "+str(value)+" shifts\n")
			three_Shifts=0
			del working[:]
			not_out.clear()
			col+=2
		else:
			col=1
			row_inc+=5


	print shift_Limit
	document.save("CODs_In.docx")


if __name__ == "__main__":
	COD_Scheduler_3000((sys.argv[1]),(sys.argv[2]),(sys.argv[3]),(sys.argv[4]),(sys.argv[5]))






		



	
	









	







